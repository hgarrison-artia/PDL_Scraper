import pandas as pd
import docx

def is_therapeutic_class(cell):
    """
    Check if a cell’s text is intended to be a therapeutic class.
    We assume that the therapeutic class cells are bold and non-italic.
    """
    bold = False
    italic = False
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.bold:
                bold = True
            if run.font.italic:
                italic = True
    return bold and italic

def is_bold(cell):
    """
    Check if any run in this cell is italicized.
    Returns True if at least one run is italic, otherwise False.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.font.bold:
                return True
    return False

doc_path = 'GA.docx'
doc = docx.Document(doc_path)
rows_list = []  # List to store rows for our final DataFrame
current_therapeutic_class = None

# Loop over each table in the document
for table in doc.tables:
    # Assuming the first two rows are header rows; iterate from the third row onward.
    for row in table.rows[2:]:
        first_cell = row.cells[0]
        cell_text = first_cell.text.strip()
        if not cell_text:
            continue

        # If the first cell's text is both bold and non-italic, treat it as a therapeutic class header.
        if is_therapeutic_class(first_cell):
            current_therapeutic_class = cell_text
        else:
            # If the product row's pdl_name is italicized, skip this row.
            if is_bold(first_cell):
                continue

            # Treat this row as a product row.
            cells = [cell.text.strip() for cell in row.cells]
            pdl_name = cells[0]
            preferred_value = cells[1] if len(cells) > 1 else ""
            status = "Preferred" if preferred_value == "P" else "Non-Preferred"
            # Append the row with pdl_name, status, and the current therapeutic class.
            rows_list.append([pdl_name, status, current_therapeutic_class])

# Create the final DataFrame with three columns.
final_df = pd.DataFrame(rows_list, columns=['pdl_name', 'status', 'therapeutic_class'])

final_df.to_csv('GA_PDL.csv', index=False)





