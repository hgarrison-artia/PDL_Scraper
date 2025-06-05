import pandas as pd
from docx import Document
from docx.shared import Pt

def get_font_size(cell):
    """Return the first explicit run font size (in points) found in this cell,
    or fall back to the paragraph's style font size."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.size:
                return run.font.size.pt
        if para.style and para.style.font.size:
            return para.style.font.size.pt
    return None

def find_two_unique_cols(table):
    """
    Look at the text of each cell in the table to figure out which two
    columns you actually want:
      - col0 is always the "first" column
      - scan col1, col2… and pick the first one whose text (ignoring
        the first three data rows) isn't identical to col0
    Returns (first_idx, second_idx).
    """
    # build a text‐only matrix
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)
    
    # Print table dimensions for debugging
    print(f"Table dimensions: {df.shape}")
    
    # first row is header
    df.columns = df.iloc[0]
    df = df[1:]
    
    # If table has only one column, return (0, 0)
    if df.shape[1] <= 1:
        print("Warning: Table has only one column")
        return 0, 0
    
    first = df.iloc[:, 0]
    for i in range(1, df.shape[1]):
        cand = df.iloc[:, i]
        # Print column comparison for debugging
        print(f"Comparing column 0 with column {i}")
        print(f"Column 0 sample: {first.iloc[3:6].tolist()}")
        print(f"Column {i} sample: {cand.iloc[3:6].tolist()}")
        
        # Check if columns are different after first 3 rows
        if not first.iloc[3:].equals(cand.iloc[3:]):
            return 0, i
    
    # If we get here, all columns are identical after first 3 rows
    print("Warning: All columns appear identical after first 3 rows")
    # Return the first two columns as a fallback
    return 0, 1

def extract_pdl_flat(doc_path):
    doc = Document(doc_path)
    records = []
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nProcessing table {table_idx + 1}")
        # reset headers for each table
        current_main = None
        current_sub  = None
        
        try:
            # figure out which two columns hold your PDL drugs
            first_idx, second_idx = find_two_unique_cols(table)
            
            # now walk each data‐row (skip the table's very first row)
            for row in table.rows[1:]:
                for col_idx, status in ((first_idx, "Preferred"), (second_idx, "Non-Preferred")):
                    cell = row.cells[col_idx]
                    txt  = cell.text.strip()
                    if not txt:
                        continue
                    
                    size = get_font_size(cell)
                    
                    # No Font size = main header
                    if size == None:
                        current_main = txt
                        current_sub  = None
                        continue
                    
                    # 10 pt = subheader
                    if size == 10:
                        current_sub = txt
                        continue
                    
                    #  8 pt = actual drug line
                    if size == 8:
                        # stitch together main + sub
                        if current_main and current_sub:
                            tc = f"{current_main}: {current_sub}"
                        elif current_main:
                            tc = current_main
                        else:
                            tc = ""
                        
                        records.append({
                            "pdl_name":          txt,
                            "status":            status,
                            "therapeutic_class": tc
                        })
                        continue
        except Exception as e:
            print(f"Error processing table {table_idx + 1}: {str(e)}")
            continue
        
    return pd.DataFrame(records)

# usage
df = extract_pdl_flat("MS.docx")

df['therapeutic_class'] = (
    df['therapeutic_class']
      # remove the asterisked phrase (including the asterisks themselves)
      .str.replace(r'\*[^*]+\*', '', regex=True)
      # collapse any extra spaces that might remain
      .str.replace(r'\s+', ' ', regex=True)
      .str.strip()
)

df['therapeutic_class'] = (
    df['therapeutic_class']
      # remove any "DUR+ " (and trailing spaces)
      .str.replace(r'DUR\+\s*', '', regex=True)
      # collapse extra spaces just in case
      .str.replace(r'\s+', ' ', regex=True)
      .str.strip()
)

df = df[['therapeutic_class','pdl_name','status']]
df.to_csv('MS_PDL.csv', index=False)