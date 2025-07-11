import pandas as pd
import openpyxl
import docx
from docx.oxml.ns import qn


def get_cell_text_without_superscripts(cell):
    """Return cell text while ignoring superscripted runs."""
    lines = []
    for paragraph in cell.paragraphs:
        parts = []
        for run in paragraph.runs:
            if run.element.xpath('.//w:vertAlign'):
                # Skip text that is formatted as superscript (e.g. "PA")
                continue
            parts.append(run.text)
        lines.append("".join(parts))
    return "\n".join(lines).strip()


def get_cell_lines_with_format(cell):
    """Return list of dictionaries with text and font sizes for each paragraph,
    ignoring superscripted runs."""
    lines = []
    for paragraph in cell.paragraphs:
        texts = []
        fonts = []
        sizes = []
        for run in paragraph.runs:
            if run.element.xpath('.//w:vertAlign'):
                # Skip superscripted text
                continue
            texts.append(run.text)
            fonts.append(run.font.name)
            sizes.append(run.font.size.pt if run.font.size else None)
        text = "".join(texts).strip()
        if text:
            lines.append({
                "text": text,
                "fonts": fonts,
                "sizes": sizes
            })
    return lines


def is_strict_calibri_size(line, size_pt):
    """Check if all runs in the line are Calibri with the given size."""
    fonts = [f for f in line["fonts"] if f]
    sizes = [s for s in line["sizes"] if s is not None]
    if not fonts or not sizes:
        return False
    return all(f.lower() == "calibri" for f in fonts) and all(abs(s - size_pt) < 0.5 for s in sizes)


def line_starts_11_ends_9(line):
    """Return True if the line begins in 11pt and ends in 9pt."""
    sizes = [s for s in line["sizes"] if s is not None]
    if not sizes:
        return False
    start = sizes[0]
    end = sizes[-1]
    return abs(start - 11) < 0.5 and abs(end - 9) < 0.5


def merge_overflow_lines(lines):
    """Merge lines where a 9pt line continues text from a previous line."""
    merged = []
    for line in lines:
        if merged and is_strict_calibri_size(line, 9) and line_starts_11_ends_9(merged[-1]):
            merged[-1]["text"] += f" {line['text']}"
            merged[-1]["fonts"].extend(line["fonts"])
            merged[-1]["sizes"].extend(line["sizes"])
        else:
            merged.append(line)
    return merged

doc_path = 'OH.docx' 
doc = docx.Document(doc_path)

class_color = '002F86'
redrow_color = 'AB2228'
subclass_color = 'FFC500'
Class = None
Subclass = None
table_data = []


for table_idx, table in enumerate(doc.tables):
    Class = None
    Subclass = None
    for row in table.rows:
        row_values = [get_cell_text_without_superscripts(cell) for cell in row.cells if get_cell_text_without_superscripts(cell)]
        row_class = None
        # Only combine if not all cell values are identical
        if row_values and not all(val == row_values[0] for val in row_values):
            for cell in row.cells:
                shd = cell._tc.xpath('.//w:shd')
                fill = shd[0].get(qn('w:fill')) if shd else None
                if fill == class_color:
                    row_class = "".join(row_values)
                    break
        else:
            for cell in row.cells:
                shd = cell._tc.xpath('.//w:shd')
                fill = shd[0].get(qn('w:fill')) if shd else None
                if fill == class_color:
                    row_class = row_values[0]
                    break
        if row_class:
            Class = row_class
            Subclass = None  # Optionally reset Subclass when a new Class is found
        for idx, cell in enumerate(row.cells):
            value = get_cell_text_without_superscripts(cell)
            shd = cell._tc.xpath('.//w:shd')
            fill = shd[0].get(qn('w:fill')) if shd else None
            if fill == subclass_color and value:
                Subclass = value
            elif fill != class_color and fill != subclass_color and idx != len(row.cells) - 1:
                lines = merge_overflow_lines(get_cell_lines_with_format(cell))
                for line in lines:
                    drug = line['text'].strip()
                    if drug and Class and Class != 'Example Category' and fill != redrow_color:
                        status = "Preferred" if idx == 0 else "Non-Preferred"
                        table_data.append([Class, Subclass, drug, status])

df = pd.DataFrame(table_data, columns=['therapeutic_class', 'Subclass', 'pdl_name', 'status'])
df['therapeutic_class'] = df['therapeutic_class'].fillna('') + ': ' + df['Subclass'].fillna('')
df = df.drop(['Subclass'], axis=1)
df = df.drop_duplicates().reset_index(drop=True)
df.to_csv('OH_PDL.csv', index=False)